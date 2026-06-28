"""
Resume building: structured content generation and .docx rendering.
ATS-compliant; red font for deficient content; inline deficiency comments.
Quality: align bullets with job requirements and yes/no concept gaps.
"""
import logging
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt

from backend.services.gemini_llm_service import get_llm_service
from backend.models.schemas.resume import ResumeStructured, ResumeSection, ResumeBullet

logger = logging.getLogger(__name__)


RESUME_BUILDER_SYSTEM = """You are an expert resume writer and ATS specialist. You generate an ATS-friendly resume structure based on job requirements, the candidate's concept checklist results (yes / a_bit / no), and the skill gap analysis. Treat "a_bit" as partial awareness: soften deficiency wording for a_bit vs "no".

Output a single JSON object with:
- "summary": string. A 2–4 sentence professional summary that reflects the candidate's strengths (concepts they said yes to) and avoids overclaiming on gaps (concepts they said no to). Use action-oriented language.
- "summary_deficiency": string or null. If the summary cannot fully align with the role due to gaps, add one short phrase here, e.g. "[Consider strengthening X before interviews]". Otherwise null.
- "sections": list of objects. Each has "heading" (exactly one of: Summary, Skills, Experience, Projects, Education, Certifications) and "bullets" (list of objects with "text", "is_deficient", "deficiency_comment").
  - Use standard headings. Summary content goes in "summary"; do not duplicate in sections.
  - For Skills: list concrete skills; mark items the candidate said "no" to (or major gaps) with is_deficient: true and deficiency_comment like "[Needs to prepare]" or "[Not yet demonstrated]".
  - For Experience/Projects: write 2–4 bullets per section. Use past tense, action verbs, and quantification where possible. Mark any bullet that overclaims on a gap with is_deficient: true and a short deficiency_comment, e.g. "[No measurable impact provided]" or "[Concept not yet demonstrated]".
  - Keep total content to fit within 2 pages. Prefer quality and relevance over length.
- "max_pages": 2.

Rules:
- Do not fabricate company names, job titles, or metrics. Use placeholder phrasing like "Delivered X" or "Improved Y" where specific numbers are not known.
- deficiency_comment must be short and in square brackets. Only set is_deficient true where the gap analysis or yes/no answers justify it.
- Output only valid JSON. No markdown, no code fences."""


def build_resume_structured(
    extracted_skills: dict[str, Any],
    evaluation_result: dict[str, Any],
    skill_gap_summary: dict[str, Any],
    user_answers: dict[str, str],
    max_pages: int = 2,
    user_baseline: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """
    Generate resume structure. Mark deficient items from skill gaps and no-answers.
    If user_baseline is provided (skills, experience_summary, baseline_resume_json), use it to personalize content.
    """
    llm = get_llm_service()
    user = (
        f"Job requirements:\n{json.dumps(extracted_skills, indent=2)}\n\n"
        f"Evaluation (yes/no concept results):\n{json.dumps(evaluation_result, indent=2)}\n\n"
        f"Skill gap analysis:\n{json.dumps(skill_gap_summary, indent=2)}\n\n"
        f"Candidate yes/no answers (for alignment):\n{json.dumps(user_answers)}"
    )
    if user_baseline:
        user += f"\n\nCandidate baseline (use real experience/skills when present):\n{json.dumps(user_baseline, indent=2)}"
    try:
        raw = llm.invoke_json_dict(RESUME_BUILDER_SYSTEM, user, stage="resume_builder")
        raw["max_pages"] = max_pages
        ResumeStructured.model_validate(raw)
        return raw
    except Exception as e:
        logger.exception("build_resume_structured failed")
        raise


def render_resume_docx(resume: ResumeStructured, path: Path) -> None:
    """
    Write resume to .docx. No tables; standard headings; red font for deficient content.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    if resume.summary:
        p = doc.add_paragraph()
        p.add_run("Summary\n").bold = True
        p.add_run(resume.summary)
        if resume.summary_deficiency:
            r = p.add_run(" " + resume.summary_deficiency)
            r.font.color.rgb = __red_rgb()
        doc.add_paragraph()

    for section in resume.sections:
        doc.add_paragraph(section.heading).runs[0].bold = True
        for bullet in section.bullets:
            p = doc.add_paragraph(style="List Bullet")
            if bullet.is_deficient and bullet.deficiency_comment:
                p.add_run(bullet.text)
                r = p.add_run(" " + bullet.deficiency_comment)
                r.font.color.rgb = __red_rgb()
            elif bullet.is_deficient:
                r = p.add_run(bullet.text)
                r.font.color.rgb = __red_rgb()
            else:
                p.add_run(bullet.text)
        doc.add_paragraph()

    doc.save(path)


def __red_rgb():
    from docx.shared import RGBColor
    return RGBColor(200, 0, 0)


def render_resume_pdf(resume: ResumeStructured, path: Path) -> None:
    """Write resume to PDF (simple text layout)."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem

    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    story = []
    if resume.summary:
        story.append(Paragraph("<b>Summary</b>", styles["Heading2"]))
        story.append(Paragraph((resume.summary or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["Normal"]))
        if resume.summary_deficiency:
            story.append(Paragraph(f'<font color="red">{resume.summary_deficiency.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")}</font>', styles["Normal"]))
        story.append(Spacer(1, 12))
    for section in resume.sections:
        story.append(Paragraph(f"<b>{section.heading.replace('&', '&amp;')}</b>", styles["Heading2"]))
        for bullet in section.bullets:
            text = bullet.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            if bullet.is_deficient and bullet.deficiency_comment:
                text += f' <font color="red">{bullet.deficiency_comment.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")}</font>'
            elif bullet.is_deficient:
                text = f'<font color="red">{text}</font>'
            story.append(Paragraph(f"• {text}", styles["Normal"]))
        story.append(Spacer(1, 8))
    doc.build(story)


def render_resume_docx_compact(resume: ResumeStructured, path: Path) -> None:
    """Same as render_resume_docx with smaller font and tighter spacing (compact template)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"
    if resume.summary:
        p = doc.add_paragraph()
        p.add_run("Summary\n").bold = True
        p.add_run(resume.summary)
        if resume.summary_deficiency:
            r = p.add_run(" " + resume.summary_deficiency)
            r.font.color.rgb = __red_rgb()
        doc.add_paragraph()
    for section in resume.sections:
        doc.add_paragraph(section.heading).runs[0].bold = True
        for bullet in section.bullets:
            p = doc.add_paragraph(style="List Bullet")
            if bullet.is_deficient and bullet.deficiency_comment:
                p.add_run(bullet.text)
                r = p.add_run(" " + bullet.deficiency_comment)
                r.font.color.rgb = __red_rgb()
            elif bullet.is_deficient:
                r = p.add_run(bullet.text)
                r.font.color.rgb = __red_rgb()
            else:
                p.add_run(bullet.text)
        doc.add_paragraph()
    doc.save(path)
